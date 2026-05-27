import pandas as pd
from docx import Document
import streamlit as st
import io
import base64

# ฟังก์ชันสำหรับอ่านไฟล์ภาพและแปลงเป็น base64
def get_base64_image(image_path):
    try:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode()
    except:
        return ""

# เตรียมข้อมูลภาพไว้ล่วงหน้า
bg_img_data = get_base64_image("background.jpg")

# 1. ตั้งค่าหน้าเว็บ
st.set_page_config(page_title="IT-SN Asset Management", page_icon="💻", layout="wide")

# 2. CSS ปรับแต่งสไตล์ Cyber Glassmorphism แบบจัดเต็ม
is_logged_in = st.session_state.get('logged_in', False)
bg_style = f"background-image: url('data:image/jpeg;base64,{bg_img_data}');" if is_logged_in else "background-color: #000000; background-image: radial-gradient(circle at 80% 20%, #1a0b2e, #000000);"

st.markdown(f"""
    <style>
    .stApp {{ 
        {bg_style}
        background-size: cover;
        background-position: center;
        background-attachment: fixed;
        background-repeat: no-repeat;
    }}

    /* สไตล์ Sidebar แบบ Cyber Glassmorphism */
    [data-testid="stSidebar"] {{
        background: rgba(0, 0, 0, 0.7) !important;
        backdrop-filter: blur(25px) !important;
        border-right: 1px solid #00ffff !important;
    }}
    
    /* หัวข้อ sidebar */
    [data-testid="stSidebar"] h2, [data-testid="stSidebar"] header {{
        color: #00ffff !important;
        text-transform: uppercase !important;
        letter-spacing: 2px !important;
        border-bottom: 1px solid #00ffff !important;
        padding-bottom: 10px !important;
        margin-bottom: 20px !important;
    }}
    
    /* ปรับแต่งปุ่ม Radio ใน Sidebar ให้มีขนาดเท่ากันเป๊ะ */
    div[role="radiogroup"] {{
        display: flex;
        flex-direction: column;
        gap: 12px;
    }}
    div[role="radiogroup"] > label {{
        background: rgba(0, 0, 0, 0.4) !important;
        border: 2px solid #00ffff !important;
        border-radius: 12px !important;
        padding: 15px 20px !important;
        width: 100% !important;
        color: #ffffff !important;
        text-align: center !important;
        transition: all 0.3s ease !important;
        font-weight: 600 !important;
    }}
    div[role="radiogroup"] > label:hover {{
        background: rgba(0, 255, 255, 0.1) !important;
        box-shadow: 0 0 15px rgba(0, 255, 255, 0.3) !important;
    }}
    div[role="radiogroup"] > label:has(input:checked) {{
        background: rgba(0, 255, 255, 0.2) !important;
        border: 2px solid #00ffff !important;
        box-shadow: 0 0 20px rgba(0, 255, 255, 0.5) !important;
    }}
    div[role="radiogroup"] input[type="radio"] {{ display: none; }}
    
    /* สไตล์ปุ่ม Action ทั่วไป (เช่น ปุ่ม SIGN IN) */
    .stButton > button {{
        background: transparent !important;
        color: #00ffff !important;
        border: 2px solid #00ffff !important;
        border-radius: 50px !important;
        padding: 10px 30px !important;
        text-transform: uppercase !important;
        letter-spacing: 2px !important;
        transition: all 0.3s ease !important;
        width: 100%;
    }}
    .stButton > button:hover {{
        background: rgba(0, 255, 255, 0.1) !important;
        box-shadow: 0 0 20px rgba(0, 255, 255, 0.5) !important;
    }}
    
    /* เจาะจงปรับแต่งปุ่มใน Sidebar ให้มีขนาดเล็กลง กะทัดรัด ไม่ยืดเต็มจอ */
    [data-testid="stSidebar"] .stButton > button {{
        padding: 5px 15px !important;
        font-size: 14px !important;
        letter-spacing: 1px !important;
        width: auto !important;
        min-width: 150px;
        border-radius: 20px !important;
        margin-bottom: 10px !important;
    }}
    
    /* ช่องกรอกข้อมูล */
    .stTextInput > div > div > input {{
        background: rgba(255, 255, 255, 0.05) !important;
        border: 1px solid #00ffff !important;
        border-radius: 10px !important;
        color: white !important;
    }}

    /* สไตล์กล่องครอบสี่เหลี่ยม Glassmorphism ตอน Login */
    [data-testid="stForm"] {{
        background: rgba(0, 0, 0, 0.6) !important;
        backdrop-filter: blur(20px) !important;
        border: 1px solid rgba(0, 255, 255, 0.3) !important;
        border-radius: 20px !important;
        padding: 40px !important;
    }}
    .login-title {{
        text-align: center;
        color: white;
        font-size: 2rem;
        font-weight: 700;
        margin-bottom: 20px;
    }}
    </style>
    """, unsafe_allow_html=True)

# 3. จัดการสถานะ
if 'logged_in' not in st.session_state: st.session_state.logged_in = False
if 'selected_ids' not in st.session_state: st.session_state.selected_ids = set()

# --- ฟังก์ชันระบบ ---
@st.cache_data # 🎯 เพิ่มตรงนี้เพื่อให้ระบบจำข้อมูลออนไลน์ไว้ ไม่ต้องดาวน์โหลดใหม่ทุกครั้งที่กดปุ่มอื่น
def get_data_from_sheets():
    SHEET_ID = "1JI5iKS6uwkvkZ1e6W5f-APAEpHwPws4gKhgpMJdzkEw"
    SHEET_NAME = "Database" 
    url = f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/gviz/tq?tqx=out:csv&sheet={SHEET_NAME}"
    df = pd.read_csv(url, encoding="utf-8", skiprows=1)
    df['unique_key'] = df.index 
    return df

def process_word_document(df_final):
    progress_text = "กำลังสร้างเอกสาร..."
    my_bar = st.progress(0, text=progress_text)
    doc = Document("template.docx")
    companies_config = [
        {"name": "บริษัท ทริปเปิล วี บรอดคาสท์ จำกัด", "code": "TVB", "table_index": 0},
        {"name": "บริษัท เทรนด์ วีจี3 จำกัด", "code": "VG3", "table_index": 1},
        {"name": "บริษัท ไทยรัฐ คอนซูเมอร์ จำกัด", "code": "TRC", "table_index": 2},
        {"name": "บริษัท ไทยรัฐ โลจิสติคส์ จำกัด", "code": "TRL", "table_index": 3},
        {"name": "บริษัท โยดาห์ บิซ จำกัด", "code": "YOD", "table_index": 4},
        {"name": "บริษัท วัชรพล จำกัด", "code": "TR", "table_index": 5},
        {"name": "บริษัท เอฟเวอร์พิงค์ จำกัด", "code": "EVP", "table_index": 6}
    ]
    summary_text = "สรุปยอดรวมอุปกรณ์ที่ส่งมอบ:\n"
    total_all_units = 0
    all_columns = df_final.columns.tolist()
    total_companies = len(companies_config)
    for i, comp in enumerate(companies_config):
        df_filtered = df_final[df_final['บริษัท'] == comp["code"]]
        my_bar.progress((i + 1) / total_companies, text=f"กำลังประมวลผล: {comp['name']}")
        table = doc.tables[comp["table_index"]]
        company_unit_count = 0
        for index, row in df_filtered.iterrows():
            new_row = table.add_row()
            company_unit_count += 1
            new_row.cells[0].text = str(len(table.rows) - 1)
            device_items = []
            item_counter = 0
            for j in range(1, 7):
                cat_col = f"ประเภท{j}"
                spec_col = "type1" if j == 1 else f"spec{j}"
                sn_col = f"Serialเครื่อง{j}"
                if spec_col in all_columns and pd.notna(row[spec_col]) and str(row[spec_col]).strip() not in ["", "-", "nan"]:
                    item_counter += 1
                    cat_prefix = f"[{str(row[cat_col]).strip()}] " if cat_col in all_columns and pd.notna(row[cat_col]) and str(row[cat_col]).strip() not in ["", "-", "nan"] else ""
                    item_text = f"- {cat_prefix}{str(row[spec_col]).strip()}"
                    if sn_col in all_columns and pd.notna(row[sn_col]) and str(row[sn_col]).strip() not in ["", "-", "nan"]:
                        item_text += f", S/N: {str(row[sn_col]).strip()}"
                    device_items.append(item_text)
            new_row.cells[1].text = "\n".join(device_items) if device_items else "-"
            dept_col = 'แผนก/ส่วนงาน' if 'แผนก/ส่วนงาน' in all_columns else ('แผนก' if 'แผนก' in all_columns else '')
            new_row.cells[2].text = str(row[dept_col]) if dept_col in all_columns and pd.notna(row[dept_col]) else "-"
            user_col = 'ชื่อ-สกุล ผู้ใช้' if 'ชื่อ-สกุล ผู้ใช้' in all_columns else ('ชื่อผู้ใช้' if 'ชื่อผู้ใช้' in all_columns else '')
            user_name = str(row[user_col]) if user_col in all_columns and pd.notna(row[user_col]) else "-"
            sec_col = 'ฝ่าย' if 'ฝ่าย' in all_columns else ''
            section = f"ฝ่าย {row[sec_col]}" if sec_col in all_columns and pd.notna(row[sec_col]) else ""
            new_row.cells[3].text = f"{section}\n({user_name})".strip()
            new_row.cells[4].text = str(item_counter) if item_counter > 0 else "-"
        summary_text += f"{comp['name']} ทั้งหมด {company_unit_count} ชุด\n"
        total_all_units += company_unit_count
    summary_text += f"รวมทั้งสิ้น {total_all_units} ชุด"
    for paragraph in doc.paragraphs:
        if "<<สรุปยอดรวม>>" in paragraph.text: paragraph.text = paragraph.text.replace("<<สรุปยอดรวม>>", summary_text)
    my_bar.empty()
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

# --- หน้าจอหลัก ---
if not st.session_state.logged_in:
    st.write("<br><br><br>", unsafe_allow_html=True) 

    col1, col2, col3 = st.columns([1.5, 2, 1.5]) 
    
    with col2:
        with st.form("login"):
            st.markdown('<div class="login-title">WELCOME</div>', unsafe_allow_html=True)
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            
            _, btn_col, _ = st.columns([1, 1, 1])
            with btn_col:
                submit_btn = st.form_submit_button("SIGN IN", use_container_width=True)
                
            if submit_btn:
                if username == "admin" and password == "1234":
                    st.session_state.logged_in = True
                    st.rerun()
                else:
                    st.error("ชื่อผู้ใช้หรือรหัสผ่านไม่ถูกต้อง")
else:
    st.title("💻 IT-SN Asset Management")
    
    # ปุ่มอัปเดตข้อมูลล่าสุด (ขนาดเล็กตาม CSS)
    if st.sidebar.button("🔄 อัปเดตข้อมูลล่าสุด"):
        st.cache_data.clear() # ล้าง Cache ข้อมูลเก่า เพื่อบังคับดึงค่าใหม่จากกูเกิลชีทออนไลน์
        st.rerun() 
        
    if st.sidebar.button("LOGOUT", key="logout_btn_unique"):
        st.session_state.logged_in = False
        st.rerun()
    
    try:
        df_full = get_data_from_sheets()
        st.sidebar.header("🔍 กรองข้อมูล")
        mode = st.sidebar.radio("เลือกโหมดค้นหา:", ["ทั้งหมด", "เลือกบริษัท", "ค้นหา"])
        df_filtered = df_full.copy()
        
        if mode == "เลือกบริษัท" and 'บริษัท' in df_filtered.columns:
            comp = st.sidebar.selectbox("เลือกบริษัท:", ["TVB", "VG3", "TRC", "TRL", "YOD", "TR", "EVP"])
            df_filtered = df_filtered[df_filtered['บริษัท'] == comp]
        elif mode == "ค้นหา":
            search = st.sidebar.text_input("กรอกคำค้นหา (ชื่อ/Serial):")
            if search:
                mask = df_filtered.apply(lambda row: row.astype(str).str.contains(search, case=False).any(), axis=1)
                df_filtered = df_filtered[mask]
        
        df_filtered['เลือกพิมพ์'] = df_filtered['unique_key'].apply(lambda x: x in st.session_state.selected_ids)
        
        edited_df = st.data_editor(
            df_filtered, 
            column_order=["เลือกพิมพ์"] + [c for c in df_filtered.columns if c != "เลือกพิมพ์"], 
            use_container_width=True, 
            hide_index=True, 
            disabled=[c for c in df_filtered.columns if c != "เลือกพิมพ์"]
        )
        
        for index, row in edited_df.iterrows():
            if row['เลือกพิมพ์']: st.session_state.selected_ids.add(row['unique_key'])
            else: st.session_state.selected_ids.discard(row['unique_key'])
            
        selected_df = df_full[df_full['unique_key'].isin(st.session_state.selected_ids)]
        
        if not selected_df.empty:
            with st.expander(f"📌 รายการที่เลือกไว้ทั้งหมด ({len(selected_df)} รายการ)"):
                st.dataframe(selected_df.drop(columns=['unique_key']), use_container_width=True, hide_index=True)
                if st.button("🗑️ ล้างรายการที่เลือกทั้งหมด"):
                    st.session_state.selected_ids = set()
                    st.rerun()
            st.download_button(label="💾 ดาวน์โหลดไฟล์ Word", data=process_word_document(selected_df), file_name="IT-SN_Form_Combined.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.info("👆 ติ๊กเลือกที่ช่อง 'เลือกพิมพ์' หน้าชื่อพนักงานที่ต้องการ")
    except Exception as e:
        st.error(f"⚠️ เกิดข้อผิดพลาด: {e}")